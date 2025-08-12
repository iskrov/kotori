import logging
from typing import Dict, Any, Optional
import tempfile
from pathlib import Path

# PDF parsing
try:
    from pypdf import PdfReader
except ImportError:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        PdfReader = None

# DOCX parsing
try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger(__name__)


class DocumentParsingError(Exception):
    """Exception raised when document parsing fails"""
    pass


class DocumentParserService:
    """Service for parsing text content from PDF and DOCX files"""

    def __init__(self):
        self.supported_types = []
        
        if PdfReader:
            self.supported_types.append('application/pdf')
            self.supported_types.append('pdf')
        
        if Document:
            self.supported_types.append('application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            self.supported_types.append('docx')

    def is_supported_type(self, content_type: str) -> bool:
        """Check if the content type is supported"""
        return content_type.lower() in self.supported_types

    def parse_file(self, file_content: bytes, content_type: str, filename: str = "") -> Dict[str, Any]:
        """
        Parse text content from uploaded file
        
        Args:
            file_content: Raw file bytes
            content_type: MIME type of the file
            filename: Original filename (for type detection fallback)
            
        Returns:
            Dictionary with parsed content and metadata
        """
        if not self.is_supported_type(content_type):
            # Try to detect type from filename
            if filename.lower().endswith('.pdf'):
                content_type = 'application/pdf'
            elif filename.lower().endswith('.docx'):
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            else:
                raise DocumentParsingError(f"Unsupported file type: {content_type}")

        try:
            if content_type.lower() in ['application/pdf', 'pdf']:
                return self._parse_pdf(file_content, filename)
            elif content_type.lower() in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'docx']:
                return self._parse_docx(file_content, filename)
            else:
                raise DocumentParsingError(f"Unsupported content type: {content_type}")
                
        except Exception as e:
            logger.error(f"Failed to parse document {filename}: {e}")
            raise DocumentParsingError(f"Document parsing failed: {e}")

    def _parse_pdf(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Parse text content from PDF file"""
        if not PdfReader:
            raise DocumentParsingError("PDF parsing not available - pypdf not installed")

        try:
            # Create temporary file for PDF processing
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file.flush()
                
                # Read PDF
                reader = PdfReader(temp_file.name)
                
                # Extract metadata
                metadata = {
                    'page_count': len(reader.pages),
                    'title': reader.metadata.title if reader.metadata else None,
                    'author': reader.metadata.author if reader.metadata else None,
                    'subject': reader.metadata.subject if reader.metadata else None,
                }
                
                # Extract text from all pages
                text_content = ""
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += f"\n--- Page {page_num + 1} ---\n"
                            text_content += page_text
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue
                
                # Clean up temp file
                Path(temp_file.name).unlink(missing_ok=True)
                
                if not text_content.strip():
                    raise DocumentParsingError("No text content found in PDF")
                
                return {
                    'text': text_content.strip(),
                    'metadata': metadata,
                    'document_type': 'pdf',
                    'filename': filename,
                    'word_count': len(text_content.split()),
                    'character_count': len(text_content)
                }
                
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            raise DocumentParsingError(f"Failed to parse PDF: {e}")

    def _parse_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Parse text content from DOCX file"""
        if not Document:
            raise DocumentParsingError("DOCX parsing not available - python-docx not installed")

        try:
            # Create temporary file for DOCX processing
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file.flush()
                
                # Read DOCX
                doc = Document(temp_file.name)
                
                # Extract metadata
                properties = doc.core_properties
                metadata = {
                    'title': properties.title,
                    'author': properties.author,
                    'subject': properties.subject,
                    'created': properties.created.isoformat() if properties.created else None,
                    'modified': properties.modified.isoformat() if properties.modified else None,
                    'paragraph_count': len(doc.paragraphs)
                }
                
                # Extract text from paragraphs
                text_content = ""
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_content += para.text + "\n"
                
                # Extract text from tables
                for table in doc.tables:
                    text_content += "\n--- Table ---\n"
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_content += " | ".join(row_text) + "\n"
                
                # Clean up temp file
                Path(temp_file.name).unlink(missing_ok=True)
                
                if not text_content.strip():
                    raise DocumentParsingError("No text content found in DOCX")
                
                return {
                    'text': text_content.strip(),
                    'metadata': metadata,
                    'document_type': 'docx',
                    'filename': filename,
                    'word_count': len(text_content.split()),
                    'character_count': len(text_content)
                }
                
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            raise DocumentParsingError(f"Failed to parse DOCX: {e}")

    def validate_file_size(self, file_content: bytes, max_size_mb: int = 10) -> bool:
        """Validate file size is within limits"""
        size_mb = len(file_content) / (1024 * 1024)
        return size_mb <= max_size_mb

    def get_supported_types_info(self) -> Dict[str, Any]:
        """Get information about supported file types"""
        return {
            'supported_types': self.supported_types,
            'pdf_available': PdfReader is not None,
            'docx_available': Document is not None,
            'max_file_size_mb': 10,
            'description': 'Supports digital PDF and DOCX files for template extraction'
        }


# Create service instance
document_parser_service = DocumentParserService()
